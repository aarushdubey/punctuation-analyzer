import streamlit as st
import pandas as pd
import re
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document

# List of punctuation keys to analyze
PUNCTUATION_KEYS = [
    "apostrophes", "colons", "commas", "curly_brackets", "double_inverted_commas",
    "ellipses", "em_dashes", "en_dashes", "exclamation_marks", "full_stops",
    "hyphens", "other_punctuation_marks", "question_marks", "round_brackets",
    "semicolons", "slashes", "square_brackets", "vertical_bars"
]

# Function to extract all text from a Word document
def extract_text(file):
    doc = Document(file)
    content = ""

    for p in doc.paragraphs:
        content += p.text + " "

    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                content += p.text + " "
        if section.footer:
            for p in section.footer.paragraphs:
                content += p.text + " "

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                content += cell.text + " "

    content = re.sub(r'(\s?\.\s?){2,}', '...', content)
    return content

# Count punctuation marks
def count_punctuation(content):
    return {
        "apostrophes": len(re.findall(r"[\'’]", content)),
        "colons": len(re.findall(r":", content)),
        "commas": len(re.findall(r",", content)),
        "curly_brackets": len(re.findall(r"\{|\}", content)),
        "double_inverted_commas": len(re.findall(r"[“”\"]", content)),
        "ellipses": len(re.findall(r"…", content)) + content.count("..."),
        "em_dashes": len(re.findall(r"—", content)),
        "en_dashes": len(re.findall(r"–", content)),
        "exclamation_marks": len(re.findall(r"!", content)),
        "full_stops": len(re.findall(r"\.", content)) - content.count("..."),
        "hyphens": len(re.findall(r"-", content)),
        "other_punctuation_marks": len(re.findall(r"[*&%$@]", content)),
        "question_marks": len(re.findall(r"\?", content)),
        "round_brackets": len(re.findall(r"\(|\)", content)),
        "semicolons": len(re.findall(r";", content)),
        "slashes": len(re.findall(r"/", content)),
        "square_brackets": len(re.findall(r"\[|\]", content)),
        "vertical_bars": len(re.findall(r"\|", content)),
    }

# Count total words
def count_words(content):
    return len(re.findall(r"\b\w+\b", content))

# Plot line graph
def plot_line_graph(df, selected_keys):
    fig, ax = plt.subplots(figsize=(12, 6))
    for key in selected_keys:
        ax.plot(df["filename"], df[key], marker='o', label=key)

    ax.set_title("Punctuation Count Comparison", fontsize=16)
    ax.set_xlabel("Document", fontsize=12)
    ax.set_ylabel("Count", fontsize=12)
    ax.legend(fontsize=10)
    ax.set_xticks(range(len(df["filename"])))
    ax.set_xticklabels(
        [label.replace(".docx", "").replace("_", " ") for label in df["filename"]],
        rotation=25, ha='right', fontsize=9, wrap=True
    )
    plt.tight_layout()
    return fig

# Plot bar graph (for single document)
def plot_bar_graph(row, selected_keys):
    fig, ax = plt.subplots(figsize=(10, 5))
    counts = [row[key] for key in selected_keys]
    ax.bar(selected_keys, counts)
    ax.set_title(f"Punctuation Count for: {row['filename']}", fontsize=16)
    ax.set_xlabel("Punctuation Type", fontsize=12)
    ax.set_ylabel("Count", fontsize=12)
    ax.tick_params(axis='x', labelrotation=45)
    plt.tight_layout()
    return fig

# Streamlit App UI
st.set_page_config(page_title="Punctuation as Grammatology")
st.title("✒️ Punctuation as Grammatology")

st.markdown("""
### 📘 How to Use
Upload `.doc` or `.docx` files. The application will generate values of different punctuation marks and visualize them in bar and line graphs.
""")

uploaded_files = st.file_uploader("Upload .docx files", type="docx", accept_multiple_files=True)

if uploaded_files:
    results = []

    for file in uploaded_files:
        with st.spinner(f"Analyzing {file.name}..."):
            text = extract_text(file)
            punctuation = count_punctuation(text)
            word_count = count_words(text)

            result = {
                "filename": file.name,
                "word_count": word_count,
                **punctuation
            }
            results.append(result)

    df = pd.DataFrame(results)
    st.success("✅ Analysis Complete!")

    st.subheader("🔍 Punctuation Summary")
    st.dataframe(df)

    csv = df.to_csv(index=False).encode("utf-8")
    st.download_button("📥 Download CSV", data=csv, file_name="punctuation_summary.csv", mime="text/csv")

    st.subheader("📈 Visualize Punctuation")
    selected = st.multiselect("Select punctuation types to plot:", options=PUNCTUATION_KEYS, default=["commas", "full_stops", "question_marks"])

    if selected:
        if len(df) == 1:
            fig = plot_bar_graph(df.iloc[0], selected)
        else:
            fig = plot_line_graph(df, selected)

        st.pyplot(fig)

        buf = BytesIO()
        fig.savefig(buf, format="png")
        st.download_button("📥 Download Graph as PNG", data=buf.getvalue(), file_name="punctuation_graph.png", mime="image/png")

# About and Team Section
st.markdown("""
---
### 📖 About
**Punctuation as Grammatology** has emerged out of curiosity about a writer's usage of punctuation in fiction.

As general discourse has it, punctuation is either disappearing or shrinking in the way written language works. For instance, the Internet is seen as a place where there is not much use for punctuation. The usage of punctuation has been shown to decrease in general. The usage of the full stop has increased in the last 200 years while the usage of the comma has decreased — a finding that suggests that writing moves towards simpler, closure-seeking or closure-affirming discourse.

While these observations and findings about writing are true at a macro level as they speak to a huge corpus of texts, it is fascinating to engage with the micro or singular. After all, the general and the particular are forever engaged in a beautiful dance with each other.

Through this web application, the experiment on one author's punctuation is being made available for others to explore punctuation in texts they wish to understand better.

*Please note that the punctuation data generated by this application is not absolute and in itself does not mean anything. It is only meant to draw attention to punctuation as grammatology, a minor site within writing that is usually understood as logocentric, and focused on words.*

---
### 👥 Team
""")

with st.expander("👩‍🏫 About Soni Wadhwa"):
    st.markdown("""
    Soni Wadhwa currently teaches Literature Studies at [SRM University, Andhra Pradesh](https://srmap.edu.in/faculty/dr-soni-wadhwa/). She works in digital humanities projects that support her larger research area of Sindhi Studies. Three of her digital projects are active. One is the [PG Sindhi Library](https://pgsindhi-library.sanchaya.net) which is a digital archive of Sindhi literature published in India. Another is the [Sindhi Halchal Archive](https://www.sindhihalchalarchive.in/about) which is dedicated to advertisements published in Sindhi books and magazines published in India. Her [Sindhi Sanchaya](https://www.sindhisanchaya.in/about) project is devoted to libraries and institutions that are home to Sindhi books and is funded by IIT Indore. She has recently received a grant from George Mason University (USA) for her research on the making of Sindhi literature as an Indian literature.
    """)

with st.expander("👩‍💻 About Leena Lokhande"):
    st.write("Information coming soon...")

with st.expander("👨‍💻 About Aarush Dubey"):
    st.write("Information coming soon...")

with st.expander("👨‍💻 About Ayush Kumar"):
    st.write("Information coming soon...")

with st.expander("👨‍💻 About Lucky Kumar"):
    st.write("Information coming soon...")

st.markdown("""
---
### 📬 Contact
For collaborations and/or feedback, write to: [wadhwa.soni@gmail.com](mailto:wadhwa.soni@gmail.com)
""")
